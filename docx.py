#!C:\Python34
# -*- coding: utf-8 -*-
#
# Date:     07/09/2017
# Author:   gutwinalex
# Version:  4
# Function: Functions related to DOCX documents

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

document = Document()


# Define the document
def create_document(template, picture):
    document = Document(template)
    document.add_picture(picture)
    return document

# add page break
def pagebreak(doc):
    document = doc
    document.add_page_break()
    return document

# add title
def add_title(doc, title):
    document = doc
    document.add_heading(title, level=0)
    return document

# add heading 1
def add_heading_l1(doc, head1):
    document = doc
    document.add_heading(head1, level=1)
    return document

# add heading 2
def add_heading_l2(doc, head2):
    document = doc
    document.add_heading(head2, level=2)
    return document

# add heading 2 number
def add_heading_l2_number(doc, head2):
    document = doc
    document.add_heading(head2, level=2, style= 'List Number')
    return document

# add heading 2
def add_heading_l2_data(doc, head2, data):
    document = doc
    para = document.add_heading(head2, level=2)
    if data != '' or data is not None:
        para.add_run(data).bold = True
    else:
        data = '------'
        para.add_run(data).bold = True
    return document

# add paragraph
def add_paragraph(doc, text):
    document = doc
    document.add_paragraph(text)
    return document

# add paragraph with bold header
def add_paragraph_with_header(doc, header, text):
    document = doc
    p = document.add_paragraph('')
    p.add_run(header).bold = True
    p.add_run(text)
    return document

# add paragraph number
def add_paragraph_number(doc, text):
    document = doc
    p = document.add_paragraph('', style='List Number')
    p.add_run(text).bold = True
    return document

def add_paragraph_bullet(doc, text):
    document = doc
    p = document.add_paragraph('', style='List Bullet')
    p.add_run(text).bold = True
    return document

# add paragraph intense quote
def add_paragraph_intense(doc, text):
    document = doc
    p = document.add_paragraph(text, style='Intense Quote')
    return document

# add paragraph with assigned data
def add_paragraph_data(doc, text, data):
    document = doc
    para = document.add_paragraph(text, style='Normal')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if data != '' or data is not None:
        para.add_run(data)
    else:
        data = '-----'
        para.add_run(data)
    return document

# add table 1_cols
def add_table_1(doc, style, cells_text_0):
    document = doc
    table = document.add_table(rows=1, cols=1, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    return document, table

# add rows to table 1_cols
def add_rows_1(doc, tab, cells_text_0):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    return document

# add table 2_cols
def add_table_2(doc, style, cells_text_0, cells_text_1):
    document = doc
    table = document.add_table(rows=1, cols=2, style=style)
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[0].width = Inches(1.5)
    hdr_cells[1].text = cells_text_1
    hdr_cells[1].width = Inches(5.5)
    return document, table

# add rows to table 2_cols
def add_rows_2(doc, tab, cells_text_0, cells_text_1):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    row_cells[0].width = Inches(1.5)
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    row_cells[1].width = Inches(5.5)
    return document

# add table 2_cols
def add_table_2_generic(doc, style, cells_text_0, cells_text_1):
    document = doc
    table = document.add_table(rows=1, cols=2, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    return document, table

# add rows to table 2_cols
def add_rows_2_generic(doc, tab, cells_text_0, cells_text_1):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    return document

# add table 2_cols
def add_table_2_simple(doc, style, cells_text_0, cells_text_1):
    document = doc
    table = document.add_table(rows=1, cols=2, style=style)
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[0].width = Inches(4.5)
    hdr_cells[1].text = cells_text_1
    hdr_cells[1].width = Inches(1)
    return document, table

# add rows to table 2_cols
def add_rows_2_simple(doc, tab, cells_text_0, cells_text_1):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    row_cells[0].width = Inches(4.5)
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    row_cells[1].width = Inches(1)
    return document

# add table 3_cols
def add_table_3(doc, style, cells_text_0, cells_text_1, cells_text_2):
    document = doc
    table = document.add_table(rows=1, cols=3, style=style)
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    #hdr_cells[0].width = Inches(1.1)
    hdr_cells[1].text = cells_text_1
    #hdr_cells[1].width = Inches(5)
    hdr_cells[2].text = cells_text_2
    #hdr_cells[2].width = Inches(1)
    return document, table

# add rows to table 3_cols
def add_rows_3(doc, tab, cells_text_0, cells_text_1, cells_text_2):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    #row_cells[0].width = Inches(1.2)
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    #row_cells[1].width = Inches(4.6)
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    #row_cells[2].width = Inches(1.2)
    return document

# add table 4_cols
def add_table_4(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3):
    document = doc
    table = document.add_table(rows=1, cols=4, style=style)
    table.autofit = False
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[0].width = Inches(1.2)
    hdr_cells[1].text = cells_text_1
    hdr_cells[1].width = Inches(1.1)
    hdr_cells[2].text = cells_text_2
    hdr_cells[2].width = Inches(3.8)
    hdr_cells[3].text = cells_text_3
    hdr_cells[3].width = Inches(1.1)
    return document, table

# add rows to table 4_cols
def add_rows_4(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    row_cells[0].width = Inches(1.2)
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    row_cells[1].width = Inches(1.1)
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    row_cells[2].width = Inches(3.8)
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    row_cells[3].width = Inches(1.1)
    return document

# add table 5_cols
def add_table_5(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3,  cells_text_4):
    document = doc
    table = document.add_table(rows=1, cols=5, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    return document, table

# add table 5_cols
def add_table_5_m(doc, style, cells_text_0, size_0, cells_text_1, size_1, cells_text_2, size_2, cells_text_3, size_3, cells_text_4, size_4):
    document = doc
    table = document.add_table(rows=1, cols=5, style=style)
    table.autofit = False
    table.columns[0].width = Inches(size_0)
    table.columns[1].width = Inches(size_1)
    table.columns[2].width = Inches(size_2)
    table.columns[3].width = Inches(size_3)
    table.columns[4].width = Inches(size_4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    return document, table

# add rows to table 5_cols
def add_rows_5(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    if cells_text_4 != None:
        try:
            row_cells[4].text = cells_text_4
        except:
            row_cells[4].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[4].text = ''
    return document

# add table 6_cols
def add_table_6(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5):
    document = doc
    table = document.add_table(rows=1, cols=6, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    hdr_cells[5].text = cells_text_5
    return document, table

# add rows to table 6_cols
def add_rows_6(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    if cells_text_4 != None:
        try:
            row_cells[4].text = cells_text_4
        except:
            row_cells[4].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[4].text = ''
    if cells_text_5 != None:
        try:
            row_cells[5].text = cells_text_5
        except:
            row_cells[5].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[5].text = ''
    return document

# add table 7_cols
def add_table_7(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6):
    document = doc
    table = document.add_table(rows=1, cols=7, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    hdr_cells[5].text = cells_text_5
    hdr_cells[6].text = cells_text_6
    return document, table

# add rows to table 7_cols
def add_rows_7(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    if cells_text_4 != None:
        try:
            row_cells[4].text = cells_text_4
        except:
            row_cells[4].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[4].text = ''
    if cells_text_5 != None:
        try:
            row_cells[5].text = cells_text_5
        except:
            row_cells[5].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[5].text = ''
    if cells_text_6 != None:
        try:
            row_cells[6].text = cells_text_6
        except:
            row_cells[6].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[6].text = ''
    return document

# add table 8_cols
def add_table_8(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6, cells_text_7):
    document = doc
    table = document.add_table(rows=1, cols=8, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    hdr_cells[5].text = cells_text_5
    hdr_cells[6].text = cells_text_6
    hdr_cells[7].text = cells_text_7
    return document, table

# add rows to table 8_cols
def add_rows_8(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6, cells_text_7):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    if cells_text_4 != None:
        try:
            row_cells[4].text = cells_text_4
        except:
            row_cells[4].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[4].text = ''
    if cells_text_5 != None:
        try:
            row_cells[5].text = cells_text_5
        except:
            row_cells[5].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[5].text = ''
    if cells_text_6 != None:
        try:
            row_cells[6].text = cells_text_6
        except:
            row_cells[6].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[6].text = ''
    if cells_text_7 != None:
        try:
            row_cells[7].text = cells_text_7
        except:
            row_cells[7].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[7].text = ''
    return document

# add table 9_cols
def add_table_9(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6, cells_text_7, cells_text_8):
    document = doc
    table = document.add_table(rows=1, cols=9, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    hdr_cells[5].text = cells_text_5
    hdr_cells[6].text = cells_text_6
    hdr_cells[7].text = cells_text_7
    hdr_cells[8].text = cells_text_8
    return document, table

# add rows to table 9_cols
def add_rows_9(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6, cells_text_7, cells_text_8):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    if cells_text_4 != None:
        try:
            row_cells[4].text = cells_text_4
        except:
            row_cells[4].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[4].text = ''
    if cells_text_5 != None:
        try:
            row_cells[5].text = cells_text_5
        except:
            row_cells[5].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[5].text = ''
    if cells_text_6 != None:
        try:
            row_cells[6].text = cells_text_6
        except:
            row_cells[6].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[6].text = ''
    if cells_text_7 != None:
        try:
            row_cells[7].text = cells_text_7
        except:
            row_cells[7].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[7].text = ''
    if cells_text_8 != None:
        try:
            row_cells[8].text = cells_text_8
        except:
            row_cells[8].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[8].text = ''
    return document

# add table 10_cols
def add_table_10(doc, style, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6, cells_text_7, cells_text_8, cells_text_9):
    document = doc
    table = document.add_table(rows=1, cols=10, style=style)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cells_text_0
    hdr_cells[1].text = cells_text_1
    hdr_cells[2].text = cells_text_2
    hdr_cells[3].text = cells_text_3
    hdr_cells[4].text = cells_text_4
    hdr_cells[5].text = cells_text_5
    hdr_cells[6].text = cells_text_6
    hdr_cells[7].text = cells_text_7
    hdr_cells[8].text = cells_text_8
    hdr_cells[9].text = cells_text_9
    return document, table

# add rows to table 10_cols
def add_rows_10(doc, tab, cells_text_0, cells_text_1, cells_text_2, cells_text_3, cells_text_4, cells_text_5, cells_text_6, cells_text_7, cells_text_8, cells_text_9):
    document = doc
    table = tab
    row_cells = table.add_row().cells
    if cells_text_0 != None:
        row_cells[0].text = cells_text_0
    else:
        row_cells[0].text = ''
    if cells_text_1 != None:
        try:
            row_cells[1].text = cells_text_1
        except:
            row_cells[1].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[1].text = ''
    if cells_text_2 != None:
        try:
            row_cells[2].text = cells_text_2
        except:
            row_cells[2].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[2].text = ''
    if cells_text_3 != None:
        try:
            row_cells[3].text = cells_text_3
        except:
            row_cells[3].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[3].text = ''
    if cells_text_4 != None:
        try:
            row_cells[4].text = cells_text_4
        except:
            row_cells[4].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[4].text = ''
    if cells_text_5 != None:
        try:
            row_cells[5].text = cells_text_5
        except:
            row_cells[5].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[5].text = ''
    if cells_text_6 != None:
        try:
            row_cells[6].text = cells_text_6
        except:
            row_cells[6].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[6].text = ''
    if cells_text_7 != None:
        try:
            row_cells[7].text = cells_text_7
        except:
            row_cells[7].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[7].text = ''
    if cells_text_8 != None:
        try:
            row_cells[8].text = cells_text_8
        except:
            row_cells[8].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[8].text = ''
    if cells_text_9 != None:
        try:
            row_cells[9].text = cells_text_9
        except:
            row_cells[9].text = 'CONTENT_ENCODING_ERROR'
    else:
        row_cells[9].text = ''
    return document

# save the document
def save_docx(doc, name):
    document = doc
    document.save(name)
